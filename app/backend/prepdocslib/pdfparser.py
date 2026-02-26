import hashlib
import html
import io
import logging
import os
import uuid
from collections.abc import AsyncGenerator
from enum import Enum
from typing import IO, Optional

import docx
import pymupdf
from azure.ai.documentintelligence.aio import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import (
    AnalyzeDocumentRequest,
    AnalyzeResult,
    DocumentFigure,
    DocumentTable,
)
from azure.core.credentials import AzureKeyCredential
from azure.core.credentials_async import AsyncTokenCredential
from azure.core.exceptions import HttpResponseError
from PIL import Image
from pypdf import PdfReader

from .page import ImageOnPage, Page
from .parser import Parser

logger = logging.getLogger("scripts")

# HybridPdfParser configurable thresholds
HYBRID_MIN_IMAGE_BYTES = 2048  # Skip images smaller than this
HYBRID_MIN_IMAGE_DIMENSION = 50  # Skip images narrower/shorter than this (px)
HYBRID_SKIP_FORMATS = {"emf", "wmf"}
HYBRID_OCR_MIN_TEXT_CHARS = 50  # Fewer chars than this with no dominant image → OCR
HYBRID_SCAN_COVERAGE_RATIO = 0.5  # Image covering more than this fraction of page → scan
HYBRID_CONTEXT_TEXT_MAX_CHARS = 2000  # Max chars for context_text field


class LocalPdfParser(Parser):
    """
    Concrete parser backed by PyPDF that can parse PDFs into pages
    To learn more, please visit https://pypi.org/project/pypdf/
    """

    async def parse(self, content: IO) -> AsyncGenerator[Page, None]:
        logger.info("Extracting text from '%s' using local PDF parser (pypdf)", content.name)

        reader = PdfReader(content)
        pages = reader.pages
        offset = 0
        for page_num, p in enumerate(pages):
            page_text = p.extract_text()
            yield Page(page_num=page_num, offset=offset, text=page_text)
            offset += len(page_text)


class LocalDocxParser(Parser):
    """
    Concrete parser backed by python-docx that can parse DOCX files into pages.
    Uses Word's lastRenderedPageBreak and explicit page break markers to split
    the document into pages matching the original Word pagination.
    """

    _HEADING_LEVELS = {
        "Heading 1": "# ",
        "Heading 2": "## ",
        "Heading 3": "### ",
        "Heading 4": "#### ",
    }
    _NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    @staticmethod
    def _has_page_break(paragraph_element) -> bool:
        """Check if a paragraph contains a page break before its text content."""
        ns = LocalDocxParser._NS
        for run in paragraph_element.findall(f"{ns}r"):
            # Explicit page break: <w:br w:type="page"/>
            for br in run.findall(f"{ns}br"):
                if br.get(f"{ns}type") == "page":
                    return True
            # Word-saved page break: <w:lastRenderedPageBreak/>
            if run.find(f"{ns}lastRenderedPageBreak") is not None:
                return True
        return False

    async def parse(self, content: IO) -> AsyncGenerator[Page, None]:
        logger.info("Extracting text from '%s' using local DOCX parser (python-docx)", content.name)

        doc = docx.Document(content)
        page_num = 0
        offset = 0
        text_parts: list[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]  # strip namespace
            if tag == "p":
                # Check for page break before this paragraph's text
                if text_parts and self._has_page_break(element):
                    page_text = "\n".join(text_parts)
                    yield Page(page_num=page_num, offset=offset, text=page_text)
                    offset += len(page_text)
                    page_num += 1
                    text_parts = []

                para = docx.text.paragraph.Paragraph(element, doc)
                style_name = para.style.name if para.style else ""
                prefix = self._HEADING_LEVELS.get(style_name, "")
                para_text = para.text.strip()
                if para_text:
                    text_parts.append(f"{prefix}{para_text}")
            elif tag == "tbl":
                table = docx.table.Table(element, doc)
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    text_parts.append(" | ".join(cells))

        # Yield the final page
        if text_parts:
            page_text = "\n".join(text_parts)
            yield Page(page_num=page_num, offset=offset, text=page_text)


class DocumentAnalysisParser(Parser):
    """
    Concrete parser backed by Azure AI Document Intelligence that can parse many document formats into pages
    To learn more, please visit https://learn.microsoft.com/azure/ai-services/document-intelligence/overview
    """

    def __init__(
        self,
        endpoint: str,
        credential: AsyncTokenCredential | AzureKeyCredential,
        model_id: str = "prebuilt-layout",
        process_figures: bool = False,
    ) -> None:
        self.model_id = model_id
        self.endpoint = endpoint
        self.credential = credential
        self.process_figures = process_figures

    async def parse(self, content: IO) -> AsyncGenerator[Page, None]:
        logger.info("Extracting text from '%s' using Azure Document Intelligence", content.name)

        async with DocumentIntelligenceClient(
            endpoint=self.endpoint, credential=self.credential
        ) as document_intelligence_client:
            # Always convert to bytes up front to avoid passing a FileStorage/stream object
            try:
                content.seek(0)
            except Exception:
                pass
            content_bytes = content.read()

            poller = None
            doc_for_pymupdf = None

            if self.process_figures:
                try:
                    poller = await document_intelligence_client.begin_analyze_document(
                        model_id="prebuilt-layout",
                        body=AnalyzeDocumentRequest(bytes_source=content_bytes),
                        output=["figures"],
                        features=["ocrHighResolution"],
                        output_content_format="markdown",
                    )
                    doc_for_pymupdf = pymupdf.open(stream=io.BytesIO(content_bytes))
                except HttpResponseError as e:
                    if e.error and e.error.code == "InvalidArgument":
                        logger.error(
                            "This document type does not support media description. Proceeding with standard analysis."
                        )
                    else:
                        logger.error(
                            "Unexpected error analyzing document for media description: %s. Proceeding with standard analysis.",
                            e,
                        )
                    poller = None

            if poller is None:
                poller = await document_intelligence_client.begin_analyze_document(
                    model_id=self.model_id,
                    body=AnalyzeDocumentRequest(bytes_source=content_bytes),
                )
            analyze_result: AnalyzeResult = await poller.result()

            offset = 0

            for page in analyze_result.pages:
                tables_on_page = [
                    table
                    for table in (analyze_result.tables or [])
                    if table.bounding_regions and table.bounding_regions[0].page_number == page.page_number
                ]
                figures_on_page = []
                if self.process_figures:
                    figures_on_page = [
                        figure
                        for figure in (analyze_result.figures or [])
                        if figure.bounding_regions and figure.bounding_regions[0].page_number == page.page_number
                    ]
                page_images: list[ImageOnPage] = []
                page_tables: list[str] = []

                class ObjectType(Enum):
                    NONE = -1
                    TABLE = 0
                    FIGURE = 1

                MaskEntry = tuple[ObjectType, Optional[int]]

                page_offset = page.spans[0].offset
                page_length = page.spans[0].length
                mask_chars: list[MaskEntry] = [(ObjectType.NONE, None)] * page_length
                # mark all positions of the table spans in the page
                for table_idx, table in enumerate(tables_on_page):
                    for span in table.spans:
                        # replace all table spans with "table_id" in table_chars array
                        for i in range(span.length):
                            idx = span.offset - page_offset + i
                            if idx >= 0 and idx < page_length:
                                mask_chars[idx] = (ObjectType.TABLE, table_idx)
                # mark all positions of the figure spans in the page
                for figure_idx, figure in enumerate(figures_on_page):
                    for span in figure.spans:
                        # replace all figure spans with "figure_id" in figure_chars array
                        for i in range(span.length):
                            idx = span.offset - page_offset + i
                            if idx >= 0 and idx < page_length:
                                mask_chars[idx] = (ObjectType.FIGURE, figure_idx)

                # build page text by replacing characters in table spans with table html
                page_text = ""
                added_objects: set[MaskEntry] = set()
                for idx, mask_char in enumerate(mask_chars):
                    object_type, object_idx = mask_char
                    if object_type == ObjectType.NONE:
                        page_text += analyze_result.content[page_offset + idx]
                    elif object_type == ObjectType.TABLE:
                        if object_idx is None:
                            raise ValueError("Expected object_idx to be set")
                        if mask_char not in added_objects:
                            table_html = DocumentAnalysisParser.table_to_html(tables_on_page[object_idx])
                            page_tables.append(table_html)
                            page_text += table_html
                            added_objects.add(mask_char)
                    elif object_type == ObjectType.FIGURE:
                        if object_idx is None:
                            raise ValueError("Expected object_idx to be set")
                        if doc_for_pymupdf is None:  # pragma: no cover
                            raise ValueError("Expected doc_for_pymupdf to be set for figure processing")
                        if mask_char not in added_objects:
                            image_on_page = await DocumentAnalysisParser.figure_to_image(
                                doc_for_pymupdf, figures_on_page[object_idx]
                            )
                            page_images.append(image_on_page)
                            page_text += image_on_page.placeholder
                            added_objects.add(mask_char)

                # We remove these comments since they are not needed and skew the page numbers
                page_text = page_text.replace("<!-- PageBreak -->", "")
                # We remove excess newlines at the beginning and end of the page
                page_text = page_text.strip()
                yield Page(
                    page_num=page.page_number - 1,
                    offset=offset,
                    text=page_text,
                    images=page_images,
                    tables=page_tables,
                )
                offset += len(page_text)

    @staticmethod
    async def figure_to_image(doc: pymupdf.Document, figure: DocumentFigure) -> ImageOnPage:
        figure_title = figure.caption.content if figure.caption and figure.caption.content else ""
        # Generate a random UUID if figure.id is None
        figure_id = figure.id or f"fig_{uuid.uuid4().hex[:8]}"
        figure_filename = f"figure{figure_id.replace('.', '_')}.png"
        logger.info("Cropping figure %s with title '%s'", figure_id, figure_title)
        placeholder = f'<figure id="{figure_id}"></figure>'
        if not figure.bounding_regions:
            return ImageOnPage(
                bytes=b"",
                page_num=0,  # 0-indexed
                figure_id=figure_id,
                bbox=(0, 0, 0, 0),
                filename=figure_filename,
                title=figure_title,
                placeholder=placeholder,
                mime_type="image/png",
            )
        if len(figure.bounding_regions) > 1:
            logger.warning("Figure %s has more than one bounding region, using the first one", figure_id)
        first_region = figure.bounding_regions[0]
        # To learn more about bounding regions, see https://aka.ms/bounding-region
        bounding_box = (
            first_region.polygon[0],  # x0 (left)
            first_region.polygon[1],  # y0 (top
            first_region.polygon[4],  # x1 (right)
            first_region.polygon[5],  # y1 (bottom)
        )
        page_number = first_region["pageNumber"]  # 1-indexed
        cropped_img, bbox_pixels = DocumentAnalysisParser.crop_image_from_pdf_page(doc, page_number - 1, bounding_box)
        return ImageOnPage(
            bytes=cropped_img,
            page_num=page_number - 1,  # Convert to 0-indexed
            figure_id=figure_id,
            bbox=bbox_pixels,
            filename=figure_filename,
            title=figure_title,
            placeholder=placeholder,
            mime_type="image/png",
        )

    @staticmethod
    def table_to_html(table: DocumentTable):
        table_html = "<figure><table>"
        rows = [
            sorted([cell for cell in table.cells if cell.row_index == i], key=lambda cell: cell.column_index)
            for i in range(table.row_count)
        ]
        for row_cells in rows:
            table_html += "<tr>"
            for cell in row_cells:
                tag = "th" if (cell.kind == "columnHeader" or cell.kind == "rowHeader") else "td"
                cell_spans = ""
                if cell.column_span is not None and cell.column_span > 1:
                    cell_spans += f" colSpan={cell.column_span}"
                if cell.row_span is not None and cell.row_span > 1:
                    cell_spans += f" rowSpan={cell.row_span}"
                table_html += f"<{tag}{cell_spans}>{html.escape(cell.content)}</{tag}>"
            table_html += "</tr>"
        table_html += "</table></figure>"
        return table_html

    @staticmethod
    def crop_image_from_pdf_page(
        doc: pymupdf.Document, page_number: int, bbox_inches: tuple[float, float, float, float]
    ) -> tuple[bytes, tuple[float, float, float, float]]:
        """
        Crops a region from a given page in a PDF and returns it as an image.

        :param pdf_path: Path to the PDF file.
        :param page_number: The page number to crop from (0-indexed).
        :param bbox_inches: A tuple of (x0, y0, x1, y1) coordinates for the bounding box, in inches.
        :return: A tuple of (image_bytes, bbox_pixels).
        """
        # Scale the bounding box to 72 DPI
        bbox_dpi = 72
        # We multiply using unpacking to ensure the resulting tuple has the correct number of elements
        x0, y0, x1, y1 = (round(x * bbox_dpi, 2) for x in bbox_inches)
        bbox_pixels = (x0, y0, x1, y1)
        rect = pymupdf.Rect(bbox_pixels)
        # Assume that the PDF has 300 DPI,
        # and use the matrix to convert between the 2 DPIs
        page_dpi = 300
        page = doc.load_page(page_number)
        pix = page.get_pixmap(matrix=pymupdf.Matrix(page_dpi / bbox_dpi, page_dpi / bbox_dpi), clip=rect)

        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        bytes_io = io.BytesIO()
        img.save(bytes_io, format="PNG")
        return bytes_io.getvalue(), bbox_pixels


class HybridPdfParser(Parser):
    """Hybrid PDF parser that triages pages as digital or scanned.

    Digital pages are processed locally with PyMuPDF (text extraction + image
    extraction with filtering).  Pages that appear to be scanned are collected
    into a sub-PDF and sent to Azure Document Intelligence, with the results
    mapped back to the original page indices.
    """

    # Map PyMuPDF image extension strings to MIME types
    _EXT_TO_MIME: dict[str, str] = {
        "png": "image/png",
        "jpeg": "image/jpeg",
        "jpg": "image/jpeg",
        "gif": "image/gif",
        "bmp": "image/bmp",
        "tiff": "image/tiff",
        "tif": "image/tiff",
    }

    def __init__(self, di_parser: Optional["DocumentAnalysisParser"] = None) -> None:
        self.di_parser = di_parser

    # ------------------------------------------------------------------
    # Page triage
    # ------------------------------------------------------------------

    def _page_needs_ocr(self, page: pymupdf.Page) -> bool:
        """Return True if *page* appears to be scanned rather than born-digital.

        Heuristics (evaluated in order):
        1. If the page has very little extractable text *and* a single image
           that covers a large fraction of the page, treat it as scanned.
        2. If there is almost no text at all (below ``HYBRID_OCR_MIN_TEXT_CHARS``)
           the page is likely scanned regardless of image layout.
        """
        text = page.get_text()
        text_len = len(text.strip())
        page_rect = page.rect
        page_area = page_rect.width * page_rect.height

        if page_area == 0:
            return text_len < HYBRID_OCR_MIN_TEXT_CHARS

        # Gather image xrefs and their rects in page-point space
        images = page.get_images(full=True)
        image_rects: list[pymupdf.Rect] = []
        for img_info in images:
            xref = img_info[0]
            try:
                rects = page.get_image_rects(xref)
                for r in rects:
                    if not r.is_empty:
                        image_rects.append(r)
            except Exception:
                logger.debug("Failed to get image rects for xref %s on page %s", xref, page.number, exc_info=True)

        # Heuristic 1: dominant full-page image with little text
        if text_len < HYBRID_OCR_MIN_TEXT_CHARS and image_rects:
            for r in image_rects:
                img_area = r.width * r.height
                if img_area / page_area > HYBRID_SCAN_COVERAGE_RATIO:
                    return True

        # Heuristic 2: almost no text at all
        if text_len < HYBRID_OCR_MIN_TEXT_CHARS:
            return True

        return False

    # ------------------------------------------------------------------
    # Image filtering helpers (mirrors officeimageextractor logic)
    # ------------------------------------------------------------------

    @staticmethod
    def _should_skip_image(image_bytes: bytes, ext: str) -> bool:
        """Return True if the extracted image should be filtered out."""
        if ext.lower() in HYBRID_SKIP_FORMATS:
            return True
        if len(image_bytes) < HYBRID_MIN_IMAGE_BYTES:
            return True
        try:
            with Image.open(io.BytesIO(image_bytes)) as img:
                w, h = img.size
                if w < HYBRID_MIN_IMAGE_DIMENSION or h < HYBRID_MIN_IMAGE_DIMENSION:
                    return True
        except Exception:
            logger.debug("Could not open image to check dimensions, skipping it", exc_info=True)
            return True
        return False

    # ------------------------------------------------------------------
    # Main parse entry-point
    # ------------------------------------------------------------------

    async def parse(self, content: IO) -> AsyncGenerator[Page, None]:
        doc_name = getattr(content, "name", "unknown")
        logger.info("Extracting text from '%s' using HybridPdfParser", doc_name)

        try:
            content.seek(0)
        except (OSError, io.UnsupportedOperation):
            pass
        content_bytes = content.read()
        doc = pymupdf.open(stream=io.BytesIO(content_bytes))

        try:
            # Phase 1: triage every page
            digital_indices: list[int] = []
            scanned_indices: list[int] = []
            for idx in range(doc.page_count):
                page = doc.load_page(idx)
                if self._page_needs_ocr(page):
                    scanned_indices.append(idx)
                else:
                    digital_indices.append(idx)

            logger.info(
                "Document '%s': %d pages local, %d pages DI",
                doc_name,
                len(digital_indices),
                len(scanned_indices),
            )

            # Phase 2a: process digital pages locally
            offset = 0
            # We'll collect all pages (digital + DI) and yield in page order
            all_pages: dict[int, Page] = {}

            seen_hashes: set[str] = set()
            img_counter = 0

            for idx in digital_indices:
                mupdf_page = doc.load_page(idx)
                page_text = mupdf_page.get_text()

                # Extract images
                page_images: list[ImageOnPage] = []
                raw_images = mupdf_page.get_images(full=True)
                for img_info in raw_images:
                    xref = img_info[0]
                    try:
                        extracted = doc.extract_image(xref)
                    except Exception:
                        logger.debug("Failed to extract image xref %s from page %d", xref, idx, exc_info=True)
                        continue
                    if not extracted or "image" not in extracted:
                        continue

                    image_bytes = extracted["image"]
                    ext = extracted.get("ext", "png")

                    if self._should_skip_image(image_bytes, ext):
                        continue

                    img_hash = hashlib.sha256(image_bytes).hexdigest()
                    if img_hash in seen_hashes:
                        continue
                    seen_hashes.add(img_hash)

                    img_counter += 1
                    figure_id = f"img_{img_counter}"
                    mime_type = self._EXT_TO_MIME.get(ext.lower(), "image/png")
                    base_name = os.path.splitext(os.path.basename(doc_name))[0]
                    img_filename = f"{base_name}_{figure_id}.{ext}"
                    placeholder = f'<figure id="{figure_id}"></figure>'

                    page_images.append(
                        ImageOnPage(
                            bytes=image_bytes,
                            page_num=idx,
                            figure_id=figure_id,
                            bbox=(0, 0, 0, 0),
                            filename=img_filename,
                            title="",
                            placeholder=placeholder,
                            mime_type=mime_type,
                            context_text=page_text[:HYBRID_CONTEXT_TEXT_MAX_CHARS],
                        )
                    )

                # Append image placeholders to page text
                for img in page_images:
                    page_text = page_text.rstrip() + "\n" + img.placeholder

                all_pages[idx] = Page(
                    page_num=idx,
                    offset=0,  # will be fixed in final pass
                    text=page_text,
                    images=page_images,
                )

            # Phase 2b: send scanned pages to DI (if any)
            if scanned_indices and self.di_parser is not None:
                # Build a sub-PDF containing only scanned pages
                new_doc = pymupdf.open()
                try:
                    for idx in scanned_indices:
                        new_doc.insert_pdf(doc, from_page=idx, to_page=idx)
                    sub_pdf_bytes = new_doc.tobytes()
                finally:
                    new_doc.close()

                sub_pdf_stream = io.BytesIO(sub_pdf_bytes)
                sub_pdf_stream.name = doc_name  # preserve original name for logging

                # Map sub-PDF page indices back to original page indices
                page_index_map: dict[int, int] = dict(enumerate(scanned_indices))

                async for di_page in self.di_parser.parse(sub_pdf_stream):
                    orig_idx = page_index_map.get(di_page.page_num, di_page.page_num)
                    all_pages[orig_idx] = Page(
                        page_num=orig_idx,
                        offset=0,  # will be fixed in final pass
                        text=di_page.text,
                        images=di_page.images,
                        tables=di_page.tables,
                    )
            elif scanned_indices:
                # No DI parser available — yield empty pages so document page count is preserved
                logger.warning(
                    "Document '%s' has %d scanned pages but no DI parser configured; these pages will have no text",
                    doc_name,
                    len(scanned_indices),
                )
                for idx in scanned_indices:
                    all_pages[idx] = Page(page_num=idx, offset=0, text="")

            # Phase 3: yield pages in original order with correct offsets
            for idx in range(doc.page_count):
                if idx in all_pages:
                    page = all_pages[idx]
                    page.offset = offset
                    yield page
                    offset += len(page.text)
        finally:
            doc.close()
